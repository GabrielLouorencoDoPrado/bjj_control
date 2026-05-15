import streamlit as st
import pandas as pd
import plotly.express as px
from datetime import datetime
from database import get_students, get_attendance
import io
import os
from docx import Document
from fpdf import FPDF

# ==========================================
# FUNÇÃO PADRÃO DE EXPORTAÇÃO (COM LOGO E SEM NAN)
# ==========================================
def exibir_opcoes_exportacao(df, base_name, title, chave_unica):
    st.markdown("---")
    st.markdown(f"### 📥 Exportar {title}")
    
    formato = st.radio(
        "Escolha o formato:",
        ["Excel (.xlsx)", "Word (.docx)", "PDF (.pdf)"],
        horizontal=True,
        key=f"radio_{chave_unica}" 
    )

    if formato == "Excel (.xlsx)":
        output = io.BytesIO()
        with pd.ExcelWriter(output, engine='openpyxl') as writer:
            df.to_excel(writer, index=False, sheet_name='Relatorio')
        st.download_button(
            label="📥 Baixar Excel", 
            data=output.getvalue(), 
            file_name=f"{base_name}.xlsx", 
            mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            key=f"btn_ex_{chave_unica}"
        )
        
    elif formato == "Word (.docx)":
        doc = Document()
        doc.add_heading(title, 0)
        tabela = doc.add_table(rows=1, cols=len(df.columns))
        tabela.style = 'Table Grid'
        hdr_cells = tabela.rows[0].cells
        for i, coluna in enumerate(df.columns):
            hdr_cells[i].text = str(coluna)
        for index, row in df.iterrows():
            row_cells = tabela.add_row().cells
            for i, valor in enumerate(row):
                # Limpa o 'nan' no Word
                val = "" if pd.isna(valor) or str(valor).lower() == 'nan' else str(valor)
                row_cells[i].text = val
        output = io.BytesIO()
        doc.save(output)
        st.download_button(
            label="📥 Baixar Word", 
            data=output.getvalue(), 
            file_name=f"{base_name}.docx", 
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            key=f"btn_wd_{chave_unica}"
        )
        
    elif formato == "PDF (.pdf)":
        pdf = FPDF()
        pdf.add_page()
        
        # --- TENTA INSERIR A LOGO (logo.png ou logo.jpg) ---
        if os.path.exists("logo.png"):
            try:
                pdf.image("logo.png", x=10, y=8, w=35)
                pdf.ln(15) 
            except: pass
        elif os.path.exists("logo.jpg"):
            try:
                pdf.image("logo.jpg", x=10, y=8, w=35)
                pdf.ln(15)
            except: pass

        # --- TÍTULO DO DOCUMENTO ---
        pdf.set_font("Arial", 'B', 16)
        safe_title = str(title).encode('latin-1', 'ignore').decode('latin-1')
        pdf.cell(0, 15, safe_title, ln=True, align='C', border='B')
        pdf.ln(8)
        
        # --- FORMATAÇÃO ESPECÍFICA PARA A LISTA DE MATRICULADOS ---
        if chave_unica == "fidelizados_limpa":
            pdf.set_font("Arial", '', 12)
            for index, row in df.iterrows():
                val = str(row.iloc[0])
                if str(val).lower() != 'nan':
                    texto_limpo = val.encode('latin-1', 'replace').decode('latin-1')
                    pdf.cell(0, 8, txt=texto_limpo, ln=True, border='B')
        
        # --- FORMATAÇÃO ESPECÍFICA PARA O DOSSIÊ DO ALUNO ---
        elif chave_unica == "dossie":
            for index, row in df.iterrows():
                # Cabeçalho da Aula
                pdf.set_font("Arial", 'B', 12)
                aula = row.get('Numero_Aula', '?')
                data = row.get('Data', '?')
                pdf.cell(0, 8, txt=f"Aula {aula} - Data: {data}", ln=True, border='B')
                
                # Corpo dos detalhes (Ignorando nan)
                pdf.set_font("Arial", '', 11)
                for col in df.columns:
                    if col not in ['Numero_Aula', 'Data']:
                        val = row[col]
                        
                        # Pula valores vazios para limpar o visual
                        if pd.isna(val) or str(val).lower() == 'nan' or str(val).strip() == '':
                            continue
                        
                        col_nome = str(col).replace('_', ' ')
                        texto = f"{col_nome}: {val}"
                        texto_limpo = str(texto).encode('latin-1', 'replace').decode('latin-1')
                        pdf.multi_cell(0, 6, txt=texto_limpo)
                pdf.ln(6)

        # --- FORMATAÇÃO PADRÃO PARA OUTROS RELATÓRIOS ---
        else:
            for index, row in df.iterrows():
                pdf.set_font("Arial", 'B', 11)
                pdf.cell(0, 8, "Registro", ln=True, border='B')
                pdf.set_font("Arial", '', 10)
                for col in df.columns:
                    val = row[col]
                    if pd.isna(val) or str(val).lower() == 'nan':
                        val = "-"
                    texto = f"{col}: {val}"
                    texto_limpo = str(texto).encode('latin-1', 'replace').decode('latin-1')
                    pdf.multi_cell(0, 6, txt=texto_limpo)
                pdf.ln(4)
            
        pdf_bytes = pdf.output(dest='S').encode('latin-1')
        st.download_button(
            label="📥 Baixar PDF", 
            data=pdf_bytes, 
            file_name=f"{base_name}.pdf", 
            mime="application/pdf",
            key=f"btn_pdf_{chave_unica}"
        )

# ==========================================
# RENDERIZAÇÃO DA TELA (INTACTA E COMPLETA)
# ==========================================
def render():
    st.header("📈 Relatórios Financeiros e Frequência")
    df_st = get_students()
    df_att = get_attendance()
    
    if df_att.empty:
        st.info("Sem dados de presença registrados para gerar relatórios.")
        return

    df_att_valido = df_att.dropna(subset=['Data']).copy()
    if df_att_valido.empty:
        st.info("Sem datas válidas registradas.")
        return

    # --- PREPARAÇÃO DOS DADOS GERAIS ---
    df_att_calc = df_att_valido.copy()
    df_att_calc = df_att_calc.sort_values(by=['Data'])
    df_att_calc['Numero_Aula'] = df_att_calc.groupby('Aluno_ID').cumcount() + 1 

    # --- CRIANDO AS TRÊS ABAS PERMANENTES ---
    tab_graficos, tab_aluno, tab_relatorio_completo = st.tabs([
        "📊 Visão Financeira", 
        "👤 Dossiê do Aluno",
        "📑 Relatório Completo"
    ])

    # ==========================================
    # ABA 1: VISÃO GERAL 
    # ==========================================
    with tab_graficos:
        meses_abrev = {
            '01': 'jan', '02': 'fev', '03': 'mar', '04': 'abr', 
            '05': 'mai', '06': 'jun', '07': 'jul', '08': 'ago', 
            '09': 'set', '10': 'out', '11': 'nov', '12': 'dez'
        }
        
        df_att_valido['Chave_Periodo'] = df_att_valido['Data'].dt.strftime('%Y-%m')
        
        def formatar_periodo(chave):
            ano, mes = chave.split('-')
            return f"{meses_abrev[mes]}/{ano[-2:]}"

        periodos_disponiveis = sorted(df_att_valido['Chave_Periodo'].unique())
        
        periodos_selecionados = st.multiselect(
            "📅 Selecione os Períodos para Análise Financeira",
            options=periodos_disponiveis,
            default=periodos_disponiveis,
            format_func=formatar_periodo
        )
        
        if not periodos_selecionados:
            st.warning("Selecione pelo menos um período na caixa acima.")
        else:
            df_st['Valor_Base'] = pd.to_numeric(df_st['Valor_Base'], errors='coerce').fillna(0)
            df_st['Desconto_Percentual'] = pd.to_numeric(df_st['Desconto_Percentual'], errors='coerce').fillna(0)
            
            dados_grafico = []
            todos_pagantes_ids = set()
            todos_leads_ids = set()

            for periodo in sorted(periodos_selecionados):
                pres_no_mes = df_att_valido[df_att_valido['Chave_Periodo'] == periodo]
                
                ids_pagantes_mes = pres_no_mes[pres_no_mes['Status_Aula'].isin(["Normal", "Reposição"])]['Aluno_ID'].unique()
                todos_pagantes_ids.update(ids_pagantes_mes)
                
                cobranca_mes = df_st[df_st['ID'].isin(ids_pagantes_mes)].copy()
                
                if not cobranca_mes.empty:
                    cobranca_mes['Valor_A_Pagar'] = cobranca_mes['Valor_Base'] * (1 - (cobranca_mes['Desconto_Percentual'] / 100))
                    cobranca_mes.loc[cobranca_mes['Perfil_Financeiro'] == 'Bolsista', 'Valor_A_Pagar'] = 0.0
                    faturamento_mes = cobranca_mes['Valor_A_Pagar'].sum()
                else:
                    faturamento_mes = 0.0
                    
                dados_grafico.append({
                    'Período': formatar_periodo(periodo), 
                    'Faturamento Previsto (R$)': faturamento_mes
                })
                
                ids_leads_mes = pres_no_mes[pres_no_mes['Status_Aula'].isin(["Experimental", "Visita"])]['Aluno_ID'].unique()
                todos_leads_ids.update(ids_leads_mes)

            df_grafico = pd.DataFrame(dados_grafico)
            faturamento_total_periodo = df_grafico['Faturamento Previsto (R$)'].sum()
            
            st.metric("Faturamento Acumulado (Períodos Selecionados)", f"R$ {faturamento_total_periodo:.2f}")
            
            if len(df_grafico) > 0:
                fig = px.line(
                    df_grafico, x='Período', y='Faturamento Previsto (R$)', text='Faturamento Previsto (R$)', 
                    markers=True, title="Evolução do Faturamento"
                )
                fig.update_traces(
                    textposition="top center", texttemplate="R$ %{text:.2f}", line_color='#2ECC71',              
                    fill='tozeroy', fillcolor='rgba(46, 204, 113, 0.2)', cliponaxis=False
                )
                max_y = df_grafico['Faturamento Previsto (R$)'].max()
                limite_superior = max_y * 1.2 if max_y > 0 else 100
                fig.update_layout(margin=dict(r=50, t=50))
                fig.update_yaxes(rangemode="tozero", range=[0, limite_superior])
                st.plotly_chart(fig, use_container_width=True)

            st.divider()

            col1, col2 = st.columns(2)
            with col1:
                st.subheader("Lista de Alunos Pagantes")
                cobranca_geral = df_st[df_st['ID'].isin(todos_pagantes_ids)].copy()
                
                if not cobranca_geral.empty:
                    cobranca_geral['Valor_Mensal'] = cobranca_geral['Valor_Base'] * (1 - (cobranca_geral['Desconto_Percentual'] / 100))
                    cobranca_geral.loc[cobranca_geral['Perfil_Financeiro'] == 'Bolsista', 'Valor_Mensal'] = 0.0
                    
                    df_pagantes_view = cobranca_geral[['Nome', 'Turma', 'Perfil_Financeiro', 'Valor_Mensal']]
                    st.dataframe(df_pagantes_view, hide_index=True, use_container_width=True)
                    
                    exibir_opcoes_exportacao(
                        df=df_pagantes_view,
                        base_name="Alunos_Pagantes",
                        title="Relatório de Pagantes",
                        chave_unica="pagantes"
                    )
                else:
                    st.info("Nenhuma mensalidade gerada.")

            with col2:
                st.subheader("Leads (Experimentais/Visitas)")
                leads_geral = df_st[df_st['ID'].isin(todos_leads_ids)]
                leads_puros = leads_geral[~leads_geral['ID'].isin(todos_pagantes_ids)]
                
                if not leads_puros.empty:
                    st.metric("Total de Prospectos", len(leads_puros))
                    for _, lead in leads_puros.iterrows():
                        tel = str(lead.get('Telefone', ''))
                        msg = f"Olá {lead['Nome']}, vi que você treinou com a gente! O que achou da aula?"
                        
                        c_nome, c_link = st.columns([3, 1])
                        c_nome.write(f"👤 {lead['Nome']} ({lead['Turma']})")
                        if tel and tel.strip() != "" and tel.lower() != "nan":
                            link_wa = f"https://wa.me/55{tel}?text={msg.replace(' ', '%20')}"
                            c_link.markdown(f"[📲 Whats]({link_wa})")
                        else:
                            c_link.write("*(Sem tel)*")
                else:
                    st.metric("Total de Prospectos", 0)
                    st.info("Nenhum lead novo neste período.")

    # ==========================================
    # ABA 2: DOSSIÊ ESPECÍFICO DO ALUNO 
    # ==========================================
    with tab_aluno:
        st.subheader("👤 Dossiê de Frequência do Aluno")
        st.markdown("Puxe a capivara completa de um aluno específico.")
        
        lista_alunos = sorted(df_att_calc['Nome'].unique())
        if lista_alunos:
            aluno_alvo = st.selectbox("Selecione o Aluno:", ["-- Escolha um Aluno --"] + lista_alunos)
            
            if aluno_alvo != "-- Escolha um Aluno --":
                df_aluno_especifico = df_att_calc[df_att_calc['Nome'] == aluno_alvo].sort_values(by='Data', ascending=False)
                
                total_aulas = len(df_aluno_especifico)
                primeira_aula = df_aluno_especifico['Data'].min().strftime('%d/%m/%Y')
                ultima_aula = df_aluno_especifico['Data'].max().strftime('%d/%m/%Y')
                
                c1, c2, c3 = st.columns(3)
                c1.metric("Total de Aulas Registradas", total_aulas)
                c2.metric("Primeira Aula", primeira_aula)
                c3.metric("Última Aula", ultima_aula)
                
                st.divider()
                st.markdown(f"**Histórico completo de {aluno_alvo}:**")
                
                df_aluno_view = df_aluno_especifico[['Numero_Aula', 'Data', 'Status_Aula', 'Alerta', 'Observacao']].copy()
                if 'Tecnicas' in df_aluno_especifico.columns:
                    df_aluno_view.insert(3, 'Tecnicas', df_aluno_especifico['Tecnicas'])
                    
                df_aluno_view['Data'] = df_aluno_view['Data'].dt.strftime('%d/%m/%Y')
                
                st.dataframe(df_aluno_view, hide_index=True, use_container_width=True)

                exibir_opcoes_exportacao(
                    df=df_aluno_view,
                    base_name=f"Dossie_{aluno_alvo}",
                    title=f"Dossiê do Aluno - {aluno_alvo}",
                    chave_unica="dossie"
                )
        else:
            st.info("Nenhum aluno com aulas registradas.")

    # ==========================================
    # ABA 3: RELATÓRIO COMPLETO E STATUS DO MÊS
    # ==========================================
    with tab_relatorio_completo:
        st.subheader("📑 Visão Geral e Engajamento dos Alunos")
        st.markdown("Lista mostrando todo o histórico do aluno e se ele está ativo no mês selecionado.")
        
        df_att_calc['Mes_Ano_Filtro'] = df_att_calc['Data'].dt.strftime('%m/%Y')
        meses_disp = sorted(df_att_calc['Mes_Ano_Filtro'].unique(), reverse=True)
        
        if meses_disp:
            mes_filtro_completo = st.selectbox("Selecione o Mês para verificar o status de presença:", meses_disp, key="sel_mes_completo")
            
            if mes_filtro_completo:
                status_normais = ["Normal", "Reposição"]
                status_exp = ["Experimental", "Visita"]
                
                def count_normal(series): return series.isin(status_normais).sum()
                def count_exp(series): return series.isin(status_exp).sum()
                
                stats_gerais = df_att_calc.groupby('Aluno_ID').agg(
                    Total_Normais=('Status_Aula', count_normal),
                    Total_Exp=('Status_Aula', count_exp)
                ).reset_index()
                
                df_mes_filtro = df_att_calc[df_att_calc['Mes_Ano_Filtro'] == mes_filtro_completo]
                stats_mes = df_mes_filtro.groupby('Aluno_ID').size().reset_index(name='Aulas_Mes')
                
                df_base_alunos = df_st[['ID', 'Nome', 'Turma']].copy()
                
                df_final = df_base_alunos.merge(stats_gerais, left_on='ID', right_on='Aluno_ID', how='left')
                df_final = df_final.merge(stats_mes, left_on='ID', right_on='Aluno_ID', how='left')
                
                df_final = df_final.fillna(0) 
                
                df_final['Status'] = df_final['Aulas_Mes'].apply(
                    lambda x: '🟢 Está frequentando as aulas' if x > 0 else '🔴 Faltou o mês todo'
                )
                
                df_view = df_final[['Nome', 'Total_Normais', 'Total_Exp', 'Aulas_Mes', 'Status']].copy()
                df_view.rename(columns={
                    'Nome': 'Nome do Aluno',
                    'Total_Normais': 'Aulas Normais (Histórico)',
                    'Total_Exp': 'Aulas Experimentais (Histórico)',
                    'Aulas_Mes': f'Aulas em {mes_filtro_completo}'
                }, inplace=True)
                
                for col in ['Aulas Normais (Histórico)', 'Aulas Experimentais (Histórico)', f'Aulas em {mes_filtro_completo}']:
                    df_view[col] = df_view[col].astype(int)
                    
                df_view = df_view.sort_values(by='Nome do Aluno')
                
                st.dataframe(df_view, hide_index=True, use_container_width=True)
                
                # --- LISTA DE MATRICULADOS (COM TURMA) ---
                st.divider()
                st.subheader("🎯 Lista de Alunos Matriculados (Fidelizados)")
                st.markdown(f"**Critério:** Histórico total >= 3 aulas e presença em **{mes_filtro_completo}**.")
                
                df_final['Total_Historico'] = df_final['Total_Normais'] + df_final['Total_Exp']
                df_fidelizados = df_final[(df_final['Total_Historico'] >= 3) & (df_final['Aulas_Mes'] > 0)].sort_values(by='Nome')

                if not df_fidelizados.empty:
                    lista_exportacao_turma = []
                    
                    for _, row in df_fidelizados.iterrows():
                        nome_formatado = f"{row['Nome']} ({row['Turma']})"
                        st.write(nome_formatado)
                        lista_exportacao_turma.append(nome_formatado)
                    
                    df_export_fid = pd.DataFrame(lista_exportacao_turma, columns=['Aluno Matriculado'])
                    exibir_opcoes_exportacao(
                        df_export_fid, 
                        f"Matriculados_{mes_filtro_completo.replace('/', '_')}", 
                        "Lista de Matriculados", 
                        "fidelizados_limpa"
                    )
                else:
                    st.info("Nenhum aluno atingiu o critério de matriculado neste mês.")

                exibir_opcoes_exportacao(
                    df=df_view,
                    base_name=f"Engajamento_Alunos_{mes_filtro_completo.replace('/', '_')}",
                    title=f"Relatório de Engajamento - {mes_filtro_completo}",
                    chave_unica="relatorio_engajamento"
                )
        else:
            st.info("Nenhuma aula registrada.")